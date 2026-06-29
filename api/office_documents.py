from __future__ import annotations

import io
from pathlib import Path

CLAIMED_OFFICE_EXTENSIONS = frozenset({".docx", ".xlsx", ".pptx"})
CLAIMED_OFFICE_FORMATS = frozenset({"docx", "xlsx", "pptx"})
OFFICE_PREVIEW_KIND = "office"
OFFICE_RENDER_MODE = "code"
OFFICE_DEPENDENCY_HINT = (
    "Office preview is not available on this server. Install python-docx, "
    "openpyxl, and python-pptx to enable it: pip install python-docx openpyxl "
    "python-pptx"
)
OFFICE_PREVIEW_TRUNCATED_NOTICE = "[Preview truncated: Office content exceeds safe limits]"
MAX_OFFICE_PREVIEW_CHARS = 120_000
MAX_DOCX_PREVIEW_BLOCKS = 2_000
MAX_DOCX_TABLE_CELLS = 5_000
MAX_XLSX_PREVIEW_SHEETS = 20
MAX_XLSX_PREVIEW_ROWS_PER_SHEET = 500
MAX_XLSX_PREVIEW_CELLS_PER_SHEET = 5_000
MAX_PPTX_PREVIEW_SLIDES = 100
MAX_PPTX_PREVIEW_SHAPES_PER_SLIDE = 200

_WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DEFAULT_DOCX_SECTION_SIGNATURE = None

_DOCX_BODY_CHILDREN = {f"{_WORD_NAMESPACE}p", f"{_WORD_NAMESPACE}sectPr"}
_DOCX_PARAGRAPH_CHILDREN = {f"{_WORD_NAMESPACE}pPr", f"{_WORD_NAMESPACE}r"}
_DOCX_SAFE_PARAGRAPH_PROPERTY_CHILDREN = {f"{_WORD_NAMESPACE}pStyle"}
_DOCX_RUN_CHILDREN = {f"{_WORD_NAMESPACE}t"}


def _office_dependency_import_error() -> ImportError:
    return ImportError(OFFICE_DEPENDENCY_HINT)


def _load_docx_document():
    try:
        from docx import Document as document_factory
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return document_factory


def _load_workbook_reader():
    try:
        from openpyxl import load_workbook as workbook_reader
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return workbook_reader


def _load_presentation_ctor():
    try:
        from pptx import Presentation as presentation_ctor
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return presentation_ctor


def is_claimed_office_path(path: str | Path) -> bool:
    return Path(str(path)).suffix.lower() in CLAIMED_OFFICE_EXTENSIONS


def _office_format_for_path(path: str | Path) -> str:
    return Path(str(path)).suffix.lower().lstrip(".")


def _normalise_preview_text(value) -> str:
    if value is None:
        return ""
    return str(value).replace("\r", "\n").replace("\n", " ").strip()


def _preview_line_count(content: str) -> int:
    if not content:
        return 1
    return content.count("\n") + 1


def _finalize_preview_text(content: str, truncated: bool = False) -> tuple[str, bool]:
    text = (content or "").strip()
    if len(text) > MAX_OFFICE_PREVIEW_CHARS:
        text = text[:MAX_OFFICE_PREVIEW_CHARS].rstrip()
        truncated = True
    if truncated:
        text = f"{text}\n\n{OFFICE_PREVIEW_TRUNCATED_NOTICE}" if text else OFFICE_PREVIEW_TRUNCATED_NOTICE
    return text, truncated


def _docx_preview_text(document) -> tuple[str, bool]:
    chunks: list[str] = []
    paragraphs_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    tables_by_element = {
        table._tbl: (table_index, table) for table_index, table in enumerate(document.tables, start=1)
    }
    body_blocks_seen = 0
    table_cells_seen = 0
    truncated = False
    for child in document._element.body:
        if child.tag == f"{_WORD_NAMESPACE}sectPr":
            continue
        body_blocks_seen += 1
        if body_blocks_seen > MAX_DOCX_PREVIEW_BLOCKS:
            truncated = True
            break
        if child.tag == f"{_WORD_NAMESPACE}p":
            paragraph = paragraphs_by_element.get(child)
            if paragraph is not None:
                chunks.append(paragraph.text or "")
            continue
        if child.tag != f"{_WORD_NAMESPACE}tbl":
            continue
        table_index, table = tables_by_element[child]
        table_lines = [f"Table {table_index}"]
        for row in table.rows:
            cells = []
            for cell in row.cells:
                table_cells_seen += 1
                if table_cells_seen > MAX_DOCX_TABLE_CELLS:
                    truncated = True
                    break
                cells.append(_normalise_preview_text(cell.text))
            if truncated:
                break
            table_lines.append("\t".join(cells))
        chunks.append("\n".join(table_lines))
        if truncated:
            break
    return _finalize_preview_text("\n".join(chunks), truncated)


def _docx_paragraph_properties_are_safe(properties) -> bool:
    for child in properties:
        if child.tag not in _DOCX_SAFE_PARAGRAPH_PROPERTY_CHILDREN:
            return False
        if child.tag == f"{_WORD_NAMESPACE}pStyle" and child.get(f"{_WORD_NAMESPACE}val") != "Normal":
            return False
    return True


def _docx_xml_signature(element) -> tuple:
    attributes = tuple(
        sorted(
            (key, value)
            for key, value in element.attrib.items()
            if not key.rsplit("}", 1)[-1].startswith("rsid")
        )
    )
    children = tuple(_docx_xml_signature(child) for child in element)
    text = (element.text or "").strip()
    return element.tag, attributes, text, children


def _default_docx_section_signature() -> tuple:
    global _DEFAULT_DOCX_SECTION_SIGNATURE
    if _DEFAULT_DOCX_SECTION_SIGNATURE is None:
        document = _load_docx_document()()
        _DEFAULT_DOCX_SECTION_SIGNATURE = tuple(
            _docx_xml_signature(child) for child in document._element.body.sectPr
        )
    return _DEFAULT_DOCX_SECTION_SIGNATURE


def _docx_section_properties_are_safe(section_properties) -> bool:
    return tuple(_docx_xml_signature(child) for child in section_properties) == _default_docx_section_signature()


def _docx_editability(document) -> tuple[bool, str | None]:
    body = document._element.body
    for child in body:
        if child.tag not in _DOCX_BODY_CHILDREN:
            return False, "docx contains unsupported structures"
        if child.tag == f"{_WORD_NAMESPACE}sectPr" and not _docx_section_properties_are_safe(child):
            return False, "docx contains unsupported section content"
    for paragraph in document.paragraphs:
        for child in paragraph._p:
            if child.tag not in _DOCX_PARAGRAPH_CHILDREN:
                return False, "docx contains unsupported paragraph structures"
            if child.tag == f"{_WORD_NAMESPACE}pPr" and not _docx_paragraph_properties_are_safe(child):
                return False, "docx contains unsupported paragraph structures"
        for run in paragraph.runs:
            for child in run._r:
                if child.tag not in _DOCX_RUN_CHILDREN:
                    return False, "docx contains unsupported inline content"
    return True, None


def _preview_docx(raw: bytes) -> tuple[str, bool, str | None, bool]:
    try:
        document = _load_docx_document()(io.BytesIO(raw))
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read DOCX preview") from exc
    content, truncated = _docx_preview_text(document)
    editable, reason = _docx_editability(document)
    if truncated and editable:
        editable = False
        reason = "docx preview exceeds safe limits"
    return content, editable, reason, truncated


def _preview_xlsx(raw: bytes) -> tuple[str, bool]:
    try:
        workbook = _load_workbook_reader()(io.BytesIO(raw), data_only=True, read_only=True)
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read XLSX preview") from exc
    chunks: list[str] = []
    truncated = False
    try:
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            if sheet_index > MAX_XLSX_PREVIEW_SHEETS:
                truncated = True
                break
            sheet_lines = [f"Sheet: {sheet.title}"]
            rows_seen = 0
            cells_seen = 0
            for row in sheet.iter_rows(values_only=True):
                rows_seen += 1
                if rows_seen > MAX_XLSX_PREVIEW_ROWS_PER_SHEET:
                    truncated = True
                    break
                values = []
                for value in row:
                    cells_seen += 1
                    if cells_seen > MAX_XLSX_PREVIEW_CELLS_PER_SHEET:
                        truncated = True
                        break
                    values.append(_normalise_preview_text(value))
                if truncated:
                    break
                if any(values):
                    sheet_lines.append("\t".join(values))
            chunks.append("\n".join(sheet_lines).strip())
            if truncated:
                break
    finally:
        close = getattr(workbook, "close", None)
        if callable(close):
            close()
    if not chunks:
        return _finalize_preview_text("Empty workbook", truncated)
    return _finalize_preview_text("\n\n".join(chunk for chunk in chunks if chunk).strip() or "Empty workbook", truncated)


def _preview_pptx(raw: bytes) -> tuple[str, bool]:
    try:
        presentation = _load_presentation_ctor()(io.BytesIO(raw))
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read PPTX preview") from exc
    chunks: list[str] = []
    truncated = False
    for slide_index, slide in enumerate(presentation.slides, start=1):
        if slide_index > MAX_PPTX_PREVIEW_SLIDES:
            truncated = True
            break
        slide_lines = [f"Slide {slide_index}"]
        shapes_seen = 0
        for shape in slide.shapes:
            shapes_seen += 1
            if shapes_seen > MAX_PPTX_PREVIEW_SHAPES_PER_SLIDE:
                truncated = True
                break
            text = _normalise_preview_text(getattr(shape, "text", ""))
            if text:
                slide_lines.append(text)
        if len(slide_lines) == 1:
            slide_lines.append("(empty slide)")
        chunks.append("\n".join(slide_lines).strip())
        if truncated:
            break
    if not chunks:
        return _finalize_preview_text("Empty presentation", truncated)
    return _finalize_preview_text(
        "\n\n".join(chunk for chunk in chunks if chunk).strip() or "Empty presentation",
        truncated,
    )


def preview_office_document(path: str | Path, raw: bytes) -> dict:
    office_format = _office_format_for_path(path)
    if office_format not in CLAIMED_OFFICE_FORMATS:
        raise ValueError(f"Unsupported Office format: {path}")

    truncated = False
    if office_format == "docx":
        content, editable, reason, truncated = _preview_docx(raw)
    elif office_format == "xlsx":
        content, truncated = _preview_xlsx(raw)
        editable, reason = False, "xlsx preview is read-only in this slice"
    elif office_format == "pptx":
        content, truncated = _preview_pptx(raw)
        editable, reason = False, "pptx preview is read-only in this slice"
    else:  # pragma: no cover - exhaustive guard
        raise ValueError(f"Unsupported Office format: {path}")

    payload = {
        "path": str(path),
        "content": content,
        "size": len(raw),
        "lines": _preview_line_count(content),
        "preview_kind": OFFICE_PREVIEW_KIND,
        "office_format": office_format,
        "render_mode": OFFICE_RENDER_MODE,
        "editable": editable,
    }
    if reason:
        payload["edit_blocked_reason"] = reason
    if truncated:
        payload["truncated"] = True
    return payload


def _docx_bytes_from_text(content: str) -> bytes:
    document = _load_docx_document()()
    body = document._element.body
    for child in list(body):
        if child.tag != f"{_WORD_NAMESPACE}sectPr":
            body.remove(child)
    text = str(content or "").replace("\r\n", "\n").replace("\r", "\n")
    for line in text.split("\n"):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def save_office_document(path: str | Path, current_bytes: bytes, content: str) -> tuple[dict, bytes]:
    office_format = _office_format_for_path(path)
    if office_format != "docx":
        raise ValueError(f"{office_format or 'office file'} is preview-only in this slice")

    current_preview = preview_office_document(path, current_bytes)
    if not current_preview.get("editable"):
        raise ValueError(current_preview.get("edit_blocked_reason") or "DOCX document is not editable")

    saved_bytes = _docx_bytes_from_text(content)
    saved_preview = preview_office_document(path, saved_bytes)
    if not saved_preview.get("editable"):
        raise ValueError(saved_preview.get("edit_blocked_reason") or "Saved DOCX is not editable")
    return saved_preview, saved_bytes
