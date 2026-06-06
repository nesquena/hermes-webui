#!/usr/bin/env python3
"""
parse_job_brief.py — 把 HR 寫的自然語言職缺需求轉成 jobs/<id>.json

工作流程：
  1. HR 在 jobs_brief/ 用純文字 / Markdown 寫需求（任何格式皆可）
  2. 跑 parse_job_brief.py 自動把 jobs_brief/*.txt 轉成 jobs/*.json
  3. 後續 run_all.sh 流程不變

用法：
  python3 parse_job_brief.py                       # 處理 jobs_brief/ 內所有檔
  python3 parse_job_brief.py 機電主任-高雄.txt      # 只處理單一檔
  python3 parse_job_brief.py --force               # 強制重解析（不管是否較新）
  python3 parse_job_brief.py --model qwen3.5:27b   # 換解析模型
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import time
from pathlib import Path


def read_brief_file(path: Path) -> str:
    """讀取 brief 檔。支援 .txt / .md / .docx。"""
    suffix = path.suffix.lower()
    if suffix in ('.txt', '.md'):
        return path.read_text(encoding='utf-8', errors='replace')
    if suffix == '.docx':
        try:
            from docx import Document
        except ImportError:
            print('  ⚠️  未安裝 python-docx，自動安裝中...')
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', '--user', '--quiet', 'python-docx'])
            from docx import Document
        doc = Document(str(path))
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        # 表格內容也撈出來（HR 可能用表格列條件）
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        lines.append(t)
        return '\n'.join(lines)
    if suffix == '.doc':
        # 舊版 Word 用 macOS 內建 textutil
        try:
            r = subprocess.run(
                ['textutil', '-convert', 'txt', '-stdout', str(path)],
                capture_output=True, text=True, timeout=30,
            )
            return r.stdout
        except Exception as e:
            print(f'  ⚠️  .doc 轉檔失敗: {e}')
            return ''
    raise ValueError(f'不支援的格式: {suffix}')

ROOT = Path(__file__).resolve().parent
JOBS_DIR = ROOT / 'jobs'
BRIEF_DIR = ROOT / 'jobs_brief'

DEFAULT_MODEL = 'gemma4:e4b'   # 從 26b 改 e4b：小模型對「文字→JSON」任務已足夠且穩定

# 預設 scoring 模板（LLM 沒明確抽到時的 fallback）
DEFAULT_SCORING_KEYWORDS = {
    '機電/水電': {
        'construction_kw': ['營造', '建設', '建築師事務所'],
        'title_kw': ['機電', '水電', '監造', '主任', '空調', '機水電', 'PCM', 'Construction Manager'],
        'public_kw': ['公共工程', '道路', '機場', '捷運', '車站', '醫院', '學校', '圖書館', '博物館'],
    },
    '工地/營造': {
        'construction_kw': ['營造', '建設', '建築師事務所'],
        'title_kw': ['監造', '工地', '營造', '現場', '監工', '營建主管', '工地主任', '副主任'],
        'public_kw': ['公共工程', '道路', '橋樑', '機場', '捷運', '車站', '醫院', '學校'],
    },
}

DEFAULT_RESIDENTIAL_KW = ['集合住宅', '住宅大樓', '社區', '集合大樓']
DEFAULT_AUTOBIO_KW = ['認真', '負責', '細心', '溝通', '協調', '解決問題', '承擔', '完成', '用心', '主動']


def call_llm(prompt: str, model: str = DEFAULT_MODEL,
             endpoint: str = 'http://localhost:11434') -> str | None:
    """呼叫 Ollama LLM，回傳 raw text。"""
    import urllib.request
    import urllib.error
    payload = {
        'model': model,
        'prompt': prompt,
        'stream': False,
        'keep_alive': '10m',         # 連續跑時模型留在記憶體 10 分鐘，不重載
        'options': {
            'temperature': 0,
            'num_predict': 4000,    # e4b 9B 模型，4000 token 已夠
            'num_ctx': 4096,        # brief 多 500 字，4096 ctx 綽綽有餘
        },
    }
    req = urllib.request.Request(
        f'{endpoint}/api/generate',
        data=json.dumps(payload).encode('utf-8'),
        headers={'Content-Type': 'application/json'},
    )
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            return json.loads(resp.read()).get('response', '').strip()
    except Exception as e:
        print(f'  ✗ LLM 呼叫失敗: {e}')
        return None


def extract_json(text: str) -> dict | None:
    """從 LLM 輸出抽出 JSON 物件（用平衡括號掃描，能處理巢狀 + ```json 包覆）"""
    if not text:
        return None
    # 直接整段試
    try:
        return json.loads(text)
    except Exception:
        pass
    # 找第一個外層平衡的 { ... }（會跳過字串內的 {}）
    in_string = False
    escape = False
    depth, start = 0, -1
    for i, ch in enumerate(text):
        if escape:
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            if depth == 0:
                start = i
            depth += 1
        elif ch == '}' and depth > 0:
            depth -= 1
            if depth == 0 and start >= 0:
                try:
                    return json.loads(text[start:i + 1])
                except Exception:
                    start = -1
    return None


def derive_job_id(filename_stem: str, display_name: str) -> str:
    """從檔名或 display_name 推 job_id。優先用檔名（HR 寫的）。"""
    stem = filename_stem.lower()
    # 中文職務 → 拼音替換規則
    role_map = {
        '機電': 'jidian', '水電': 'shuidian', '工地': 'gongdi',
        '營造': 'yingzao', '監造': 'jiandao', '建築': 'jianzhu',
        '採購': 'caigou', '品管': 'pinguan',
        '成控': 'chengkong', '估算': 'gusuan', '跑照': 'paozhao',
        '秘書': 'mishu', '主管': 'zhuguan', '高階': 'gaojie',
        '會計': 'kuaiji', '人資': 'renzi', '行政': 'xingzheng',
        '業務': 'yewu', '設計': 'sheji', '安衛': 'anwei',
        '工程師': 'gongchengshi', '專員': 'zhuanyuan',
    }
    city_map = {
        '台北': 'taipei', '新北': 'xinbei', '桃園': 'taoyuan', '新竹': 'hsinchu',
        '台中': 'taichung', '彰化': 'changhua', '南投': 'nantou',
        '雲林': 'yunlin', '嘉義': 'chiayi', '台南': 'tainan',
        '高雄': 'kaohsiung', '屏東': 'pingtung',
    }
    role, city = '', ''
    # 優先用檔名匹配（HR 命名較精確），再 fallback 到 display_name
    for source in [filename_stem, display_name]:
        if not role:
            for k, v in role_map.items():
                if k in source:
                    role = v
                    break
        if not city:
            for k, v in city_map.items():
                if k in source:
                    city = v
                    break
    if role and city:
        return f'{role}-{city}'
    if role:
        return role
    if city:
        return city
    # fallback：把檔名清成 ascii-safe，去重複 dash、頭尾 dash
    fallback = re.sub(r'-+', '-', re.sub(r'[^a-z0-9-]', '-', stem)).strip('-')
    return fallback or 'job'


def detect_role_kw_set(brief_text: str) -> dict:
    """從文字判斷該用「機電」or「工地」的預設關鍵字。"""
    if any(k in brief_text for k in ['機電', '水電', '空調', '配線']):
        return DEFAULT_SCORING_KEYWORDS['機電/水電']
    if any(k in brief_text for k in ['工地', '工務', '營造工程師', '營建']):
        return DEFAULT_SCORING_KEYWORDS['工地/營造']
    return DEFAULT_SCORING_KEYWORDS['機電/水電']


def parse_brief(brief_text: str, filename_stem: str, model: str) -> dict | None:
    """用 LLM 把自然語言職缺需求轉成 JSON。"""
    prompt = f"""你是負責把人資的自然語言徵才需求轉成結構化 JSON 設定檔的助理。

請仔細閱讀下方需求，抽取以下欄位並輸出**純 JSON**（不要其他文字）：

{{
  "display_name": "<從第一行抽出，例如『機電(副)主任(高雄)』>",
  "search": {{
    "keyword": "<關鍵字，例如『機電 水電』>",
    "job_categories": ["<希望職類陣列，每項用 104 後台的標準名稱>"],
    "work_locations": ["<例如『高雄市』>"],
    "home_locations": ["<例如『高雄市』>"],
    "last_action_days": "<7天內 / 14天內 / 1個月內 等>",
    "work_exp_years": <整數年資，例如 3>,
    "work_exp_range": "以上",
    "majors": ["<科系陣列，例如『電機電子工程相關』>"],
    "age_min": <整數>,
    "age_max": <整數>,
    "tools": ["<工具陣列，例如『AutoCAD』>"],
    "tools_min_match": <要求至少幾項工具，預設 1>,
    "certificates": ["<證照陣列>"]
  }},
  "scoring_meta": {{
    "construction_company_keywords": ["<能加分的公司類型，例如『營造』『建設』>"],
    "title_keywords": ["<能加分的職稱關鍵字，例如『機電』『水電』『監造』『主任』>"],
    "preferred_project_types": ["<優先建案類型，例如『集合住宅』『住宅大樓』>"],
    "secondary_project_types": ["<次要建案類型，例如『公共工程』『道路』>"],
    "min_long_tenure_count": <要求至少幾份工作 ≥3 年，預設 1>,
    "level_senior_years": <主任級需要的年數，預設 5>,
    "level_junior_years": <副主任級需要的年數，預設 3>,
    "threshold": <匹配度門檻百分比整數，例如 65 或 80>
  }}
}}

注意：
- 「希望職類」如「機電工程師」要保留原文，不要改寫
- 「年齡 38~58 歲」 → age_min=38, age_max=58
- 「總年資 3 年以上」 → work_exp_years=3, work_exp_range="以上"
- 「最近活動日 14 天內」 → "14天內"（含「天內」二字）
- 工作地、居住地若只寫「高雄」，加上「市」變成「高雄市」
- 「請提供 10 份匹配度達 65% 以上」→ threshold=65
- 「至少一份工作待滿三年以上」 → min_long_tenure_count=1
- 「兩份工作待滿三年以上」 → min_long_tenure_count=2
- 工具 AutoCad → AutoCAD（大寫 CAD）
- 集合住宅、住宅大樓、社區 → preferred_project_types
- 公共工程、道路、機場 → secondary_project_types

人資需求原文：
---
{brief_text}
---

只輸出 JSON，不要解釋。"""

    print(f'  🤖 用 {model} 解析中（首次冷啟可能 60–120s）...', flush=True)
    t0 = time.time()
    raw = call_llm(prompt, model=model)
    elapsed = time.time() - t0
    if not raw:
        print(f'  ✗ {elapsed:.0f}s, LLM 無回應')
        return None
    print(f'  ✓ {elapsed:.0f}s, 解析中...')
    data = extract_json(raw)
    if not data:
        print(f'  ✗ JSON 解析失敗，raw[:300]={raw[:300]!r}')
        return None
    return data


def normalize_to_job_profile(parsed: dict, filename_stem: str, brief_text: str) -> dict:
    """把 LLM 抽出的精簡結構展開成完整 jobs/<id>.json 格式（補預設值）"""
    display_name = parsed.get('display_name', filename_stem)
    job_id = derive_job_id(filename_stem, display_name)
    role_kw = detect_role_kw_set(brief_text)

    search = parsed.get('search', {})
    meta = parsed.get('scoring_meta', {})

    # 居住地關鍵字（用於加分）
    residence_kw = ''
    locs = search.get('work_locations', []) or search.get('home_locations', [])
    if locs:
        residence_kw = locs[0].replace('市', '').replace('縣', '')

    # 等級規則
    senior_years = int(meta.get('level_senior_years', 5))
    junior_years = int(meta.get('level_junior_years', 3))
    min_long = int(meta.get('min_long_tenure_count', 1))

    profile = {
        'job_id': job_id,
        'display_name': display_name,
        '_source_brief': '',  # 由呼叫端填入
        'search': {
            'keyword': search.get('keyword', ''),
            'job_categories': search.get('job_categories', []),
            'work_locations': search.get('work_locations', []),
            'home_locations': search.get('home_locations', search.get('work_locations', [])),
            'last_action_days': search.get('last_action_days', '14天內'),
            'work_exp_years': int(search.get('work_exp_years', 3)),
            'work_exp_range': search.get('work_exp_range', '以上'),
            'majors': search.get('majors', []),
            'age_min': int(search.get('age_min', 30)),
            'age_max': int(search.get('age_max', 60)),
            'tools': search.get('tools', []),
            'tools_min_match': int(search.get('tools_min_match', 1)),
            'certificates': search.get('certificates', []),
        },
        'scoring': {
            'construction_kw': meta.get('construction_company_keywords') or role_kw['construction_kw'],
            'title_kw': meta.get('title_keywords') or role_kw['title_kw'],
            'residential_kw': meta.get('preferred_project_types') or DEFAULT_RESIDENTIAL_KW,
            'public_kw': meta.get('secondary_project_types') or role_kw['public_kw'],
            'autobio_positive': DEFAULT_AUTOBIO_KW,
            'threshold': int(meta.get('threshold', 65)),
            'level_rules': {
                'senior': {'min_years': senior_years, 'min_projects': 2, 'name': '主任級'},
                'junior': {'min_years': junior_years, 'min_projects': 1, 'name': '副主任級'},
            },
            'long_tenure_months': 36,
            'require_two_long_tenures': min_long >= 2,
            'residence_bonus_keyword': residence_kw,
        },
        'forward': {
            'format': 'complete',
            'summary_max_chars': 980,
            'emails': ['i00788@fong-yi.com.tw', 'fongchien19@gmail.com', '990409@fong-yi.com.tw'],
        },
        'llm': {
            'enabled': True,
            'rule_weight': 0.4,
            'llm_weight': 0.6,
            'scoring_model': 'gemma4:e4b',
            'summary_model': 'gemma4:e4b',
        },
    }
    return profile


def needs_update(brief_path: Path, json_path: Path) -> bool:
    if not json_path.exists():
        return True
    return brief_path.stat().st_mtime > json_path.stat().st_mtime


def process_one(brief_path: Path, model: str, force: bool) -> bool:
    text = read_brief_file(brief_path)
    if not text.strip():
        print(f'\n⚠️  {brief_path.name} 讀取後為空，跳過')
        return False
    print(f'\n▶ {brief_path.name}（{len(text)} 字{f" / {brief_path.suffix} 格式" if brief_path.suffix != ".txt" else ""}）')

    # 早期跳過：用 _source_brief 反查是否已解析過且 brief 沒更新
    if not force:
        for json_path in JOBS_DIR.glob('*.json'):
            try:
                existing = json.loads(json_path.read_text())
            except Exception:
                continue
            if existing.get('_source_brief') == brief_path.name and not needs_update(brief_path, json_path):
                print(f'  ⏭ 已解析為 {json_path.name}（用 --force 可重新解析）')
                return True

    parsed = parse_brief(text, brief_path.stem, model)
    if not parsed:
        return False
    profile = normalize_to_job_profile(parsed, brief_path.stem, text)
    profile['_source_brief'] = brief_path.name  # 記錄來源，方便 archive 偵測

    # 寫到 jobs/<id>.json
    out_path = JOBS_DIR / f'{profile["job_id"]}.json'

    out_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2))
    s = profile['search']
    sc = profile['scoring']
    print(f'  ✓ {out_path.name}')
    print(f'    地點={s["work_locations"]} 年齡={s["age_min"]}-{s["age_max"]} '
          f'年資≥{s["work_exp_years"]}年 工具={s["tools"]} threshold={sc["threshold"]}')
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('targets', nargs='*', help='指定特定檔案（jobs_brief/ 內），預設處理全部')
    ap.add_argument('--force', action='store_true', help='強制重解析（不管 mtime）')
    ap.add_argument('--model', default=DEFAULT_MODEL, help=f'LLM 模型（預設 {DEFAULT_MODEL}）')
    args = ap.parse_args()

    BRIEF_DIR.mkdir(parents=True, exist_ok=True)
    JOBS_DIR.mkdir(parents=True, exist_ok=True)

    if args.targets:
        files = []
        for t in args.targets:
            p = BRIEF_DIR / t
            if not p.exists():
                p = BRIEF_DIR / (t if t.endswith('.txt') else t + '.txt')
            if not p.exists():
                print(f'⚠️  找不到 {t}，跳過')
                continue
            files.append(p)
    else:
        files = sorted(
            list(BRIEF_DIR.glob('*.txt'))
            + list(BRIEF_DIR.glob('*.md'))
            + list(BRIEF_DIR.glob('*.docx'))
            + list(BRIEF_DIR.glob('*.doc'))
        )
        # 略過 _TEMPLATE 等開頭底線、以及 Word 暫存檔（~$xxx.docx）
        files = [f for f in files if not f.name.startswith('_') and not f.name.startswith('~$')]

    if not files:
        print('（沒有可處理的檔案）')
        sys.exit(0)

    print(f'共 {len(files)} 份職缺需求待處理')
    ok = sum(process_one(f, args.model, args.force) for f in files)
    print(f'\n=========================================')
    print(f'✓ 完成 {ok}/{len(files)}')


if __name__ == '__main__':
    main()
